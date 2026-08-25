# -*- coding: utf-8 -*-
"""
Gera o COMO-CONECTAR-IA.docx a partir deste script.

Existe como script (e nao como .docx solto no repositorio) para que o documento
possa ser regerado sempre que as instrucoes mudarem, sem edicao manual no Word:
    python scripts/gerar-doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Paleta real do site azzi.digital, para o documento ter a mesma identidade.
INK = RGBColor(0x23, 0x23, 0x23)
LIME = "D4EE2A"
PAPER = "F3F3F3"
ROSE = RGBColor(0xCC, 0x38, 0x58)
CINZA = RGBColor(0x6A, 0x6A, 0x6A)


def sombrear(celula, cor_hex):
    """python-docx nao expoe cor de fundo de celula; e preciso injetar o XML."""
    tc = celula._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), cor_hex)
    tc.append(shd)


def faixa(doc, texto, cor=LIME):
    """Bloco de destaque em faixa colorida, como os avisos da interface."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    sombrear(c, cor)
    p = c.paragraphs[0]
    r = p.add_run(texto)
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK
    r.bold = True
    doc.add_paragraph()
    return t


def codigo(doc, linhas):
    """Bloco de comando: monoespacado sobre fundo claro."""
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    sombrear(c, PAPER)
    p = c.paragraphs[0]
    for i, l in enumerate(linhas):
        r = p.add_run(("\n" if i else "") + l)
        r.font.name = "Consolas"
        r.font.size = Pt(10)
        r.font.color.rgb = INK
    doc.add_paragraph()


def titulo(doc, texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    for r in h.runs:
        r.font.color.rgb = INK
        r.font.name = "Poppins"
    return h


def paragrafo(doc, texto, tamanho=10.5, cor=None, italico=False, negrito=False):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.size = Pt(tamanho)
    r.font.name = "Poppins"
    r.italic = italico
    r.bold = negrito
    r.font.color.rgb = cor or INK
    return p


def construir():
    doc = Document()

    estilo = doc.styles["Normal"]
    estilo.font.name = "Poppins"
    estilo.font.size = Pt(10.5)

    for s in doc.sections:
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        s.top_margin = Inches(0.9)
        s.bottom_margin = Inches(0.9)

    # ---------------------------------------------------------------- capa
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = cap.add_run("Azzi")
    r.font.size = Pt(34)
    r.bold = True
    r.font.name = "Playfair Display"
    r.font.color.rgb = INK

    sub = doc.add_paragraph()
    r = sub.add_run("AzziBrain — como conectar a IA")
    r.font.size = Pt(19)
    r.font.name = "Poppins"
    r.font.color.rgb = INK

    paragrafo(
        doc,
        "Este guia mostra como ligar o AzziBrain a uma inteligência artificial na sua "
        "máquina. Você pode usar o Claude ou o Gemini — escolha só um caminho, o que "
        "for mais fácil para você. Cada opção leva menos de dois minutos.",
        cor=CINZA,
    )
    doc.add_paragraph()

    faixa(
        doc,
        "O AzziBrain funciona mesmo sem IA nenhuma. Sem conexão, ele responde com o "
        "conteúdo próprio já embutido. Ligar uma IA serve para gerar ideias novas, "
        "inéditas, em vez de girar entre as que já existem.",
    )

    # ------------------------------------------------------- como funciona
    titulo(doc, "Como o AzziBrain escolhe a IA sozinho", 1)
    paragrafo(
        doc,
        "Toda vez que o servidor inicia, ele procura sozinho por uma IA nesta máquina e "
        "se liga na primeira que encontrar, nesta ordem:",
    )

    t = doc.add_table(rows=6, cols=3)
    t.style = "Table Grid"
    cabecalho = ("Ordem", "O que ele procura", "Precisa de chave paga?")
    for i, txt in enumerate(cabecalho):
        c = t.cell(0, i)
        sombrear(c, "232323")
        run = c.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xF3, 0xF3, 0xF3)

    linhas = [
        ("1º", "Claude Code instalado no computador", "Não — usa seu próprio login"),
        ("2º", "Chave ANTHROPIC_API_KEY no arquivo .env", "Sim"),
        ("3º", "Gemini CLI instalado no computador", "Não — usa seu próprio login"),
        ("4º", "Chave GEMINI_API_KEY no arquivo .env", "Sim"),
        ("5º", "Nenhuma das anteriores: modo próprio", "Não"),
    ]
    for li, (a, b, c_) in enumerate(linhas, start=1):
        for ci, txt in enumerate((a, b, c_)):
            cel = t.cell(li, ci)
            run = cel.paragraphs[0].add_run(txt)
            run.font.size = Pt(10)
            run.font.color.rgb = INK
            if li == 5:
                sombrear(cel, PAPER)
    doc.add_paragraph()

    paragrafo(
        doc,
        "Para saber onde você está agora: abra http://localhost:3000 e olhe o selo no "
        "canto superior direito da página. Ele diz exatamente qual IA está ligada. "
        "O mesmo aparece na janela preta do terminal quando o servidor sobe.",
    )
    doc.add_paragraph()

    # ------------------------------------------------- opção 1: claude cli
    titulo(doc, "Opção 1 — Claude Code (recomendada)", 1)
    paragrafo(
        doc,
        "É a opção mais simples: não exige chave de API nem cartão de crédito. "
        "O AzziBrain usa o Claude que já está logado na sua máquina.",
        cor=CINZA,
        italico=True,
    )

    titulo(doc, "Passo 1 — instalar o Claude Code", 2)
    paragrafo(doc, "Abra o Prompt de Comando (ou PowerShell) e rode:")
    codigo(doc, ["npm install -g @anthropic-ai/claude-code"])

    titulo(doc, "Passo 2 — fazer login", 2)
    paragrafo(doc, "Ainda no terminal, rode o comando abaixo e siga o login no navegador:")
    codigo(doc, ["claude"])
    paragrafo(
        doc,
        "Depois que entrar, digite /exit para sair. O login fica salvo na máquina.",
        cor=CINZA,
    )

    titulo(doc, "Passo 3 — conferir se ficou disponível", 2)
    codigo(doc, ["claude --version"])
    paragrafo(
        doc,
        "Se aparecer um número de versão, está pronto. Se aparecer 'comando não "
        "encontrado', feche e abra o terminal de novo — o Windows só enxerga programas "
        "novos depois disso.",
        cor=CINZA,
    )

    titulo(doc, "Passo 4 — reiniciar o AzziBrain", 2)
    codigo(doc, ["npm run dev"])
    paragrafo(
        doc,
        "No terminal deve aparecer: IA: [OK] Claude Code (CLI local). "
        "Pronto — as ideias agora são geradas pelo Claude.",
        negrito=True,
    )
    doc.add_paragraph()

    # ------------------------------------------------ opção 2: chave claude
    titulo(doc, "Opção 2 — Claude por chave de API", 1)
    paragrafo(
        doc,
        "Use esta se você já tem uma chave da Anthropic e prefere não instalar nada.",
        cor=CINZA,
        italico=True,
    )
    paragrafo(
        doc,
        "Pegue sua chave em console.anthropic.com, crie um arquivo chamado .env na "
        "pasta do projeto e escreva dentro dele:",
    )
    codigo(
        doc,
        [
            "ANTHROPIC_API_KEY=sk-ant-sua-chave-aqui",
            "",
            "# opcional — se quiser escolher o modelo:",
            "ANTHROPIC_MODEL=claude-sonnet-5",
        ],
    )
    paragrafo(doc, "Salve o arquivo e reinicie com npm run dev.")
    doc.add_paragraph()

    # ------------------------------------------------------ opção 3: gemini
    titulo(doc, "Opção 3 — Gemini (era o original deste projeto)", 1)
    paragrafo(
        doc,
        "O AzziBrain nasceu usando o Gemini. Esse caminho continua funcionando igual, "
        "sem nenhuma perda.",
        cor=CINZA,
        italico=True,
    )

    titulo(doc, "Pelo Gemini CLI (sem chave paga)", 2)
    codigo(doc, ["npm install -g @google/gemini-cli", "gemini"])
    paragrafo(doc, "Faça o login no navegador, saia e reinicie o AzziBrain.", cor=CINZA)

    titulo(doc, "Ou por chave de API", 2)
    paragrafo(
        doc,
        "Pegue a chave em aistudio.google.com/apikey e coloque no arquivo .env:",
    )
    codigo(
        doc,
        [
            "GEMINI_API_KEY=sua-chave-aqui",
            "",
            "# opcional:",
            "GEMINI_MODEL=gemini-3.7-flash",
        ],
    )
    doc.add_paragraph()

    # ------------------------------------------------------------ forçar
    titulo(doc, "Forçar uma IA específica", 1)
    paragrafo(
        doc,
        "Se você tem mais de uma IA instalada e quer escolher qual usar, escreva no "
        ".env uma destas linhas:",
    )
    codigo(
        doc,
        [
            "AI_PROVIDER=claude-cli    # Claude instalado na máquina",
            "AI_PROVIDER=claude-api    # Claude por chave",
            "AI_PROVIDER=gemini-cli    # Gemini instalado na máquina",
            "AI_PROVIDER=gemini-api    # Gemini por chave",
            "AI_PROVIDER=heuristic     # nenhuma IA, só o conteúdo próprio",
        ],
    )

    # ------------------------------------------------------ problemas
    titulo(doc, "Se algo não funcionar", 1)

    problemas = [
        (
            "O selo continua dizendo 'Modo heurístico'",
            "O servidor só procura IA quando inicia. Pare com Ctrl+C e rode npm run dev "
            "de novo. Se ainda assim não achar, clique na setinha de recarregar dentro "
            "do próprio selo, na página.",
        ),
        (
            "'claude' não é reconhecido como comando",
            "Feche o terminal e abra outro. Programas recém-instalados só aparecem em "
            "terminais abertos depois da instalação.",
        ),
        (
            "A porta 3000 já está em uso",
            "Outro programa está usando a porta. Feche-o, ou mude o número da porta na "
            "linha 'const PORT = 3000' do arquivo server.ts.",
        ),
        (
            "A resposta demora muito",
            "Normal na primeira chamada por IA local: ela precisa iniciar. As seguintes "
            "são bem mais rápidas.",
        ),
        (
            "Aparece 'resposta determinística' no consultor",
            "Significa que a IA não respondeu e o sistema usou o conteúdo próprio para "
            "não travar. Confira o selo de IA e o terminal.",
        ),
    ]
    for tit, txt in problemas:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(tit + " — ")
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = ROSE
        r2 = p.add_run(txt)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = INK

    doc.add_paragraph()
    faixa(
        doc,
        "Privacidade: tudo roda na sua máquina. O AzziBrain não envia nada para "
        "servidores da Azzi. Se você ligar uma IA, apenas o texto do pedido vai para o "
        "provedor escolhido (Anthropic ou Google), como em qualquer uso normal deles.",
        cor="E3AEBD",
    )

    rodape = doc.add_paragraph()
    r = rodape.add_run("AzziBrain · azzi.digital · documento gerado por scripts/gerar-doc.py")
    r.font.size = Pt(8.5)
    r.font.color.rgb = CINZA
    r.italic = True

    destino = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "COMO-CONECTAR-IA.docx",
    )
    doc.save(destino)
    print("gerado:", destino)


if __name__ == "__main__":
    construir()
